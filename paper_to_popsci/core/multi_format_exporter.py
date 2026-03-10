"""
多格式导出器 - 支持 HTML、PDF、Word、Markdown
确保图片在所有格式中正常显示
"""
import base64
import re
from pathlib import Path
from typing import List, Optional, Dict, Any
from io import BytesIO
from .config import Config, normalize_path
from .logger import logger


class MultiFormatExporter:
    """多格式导出器"""

    def __init__(self):
        self.style = Config.STYLE

    def export(self, article_sections, paper_content, output_dir: Path, formats: List[str] = None) -> Dict[str, Path]:
        """
        导出多种格式

        Args:
            article_sections: 文章章节列表
            paper_content: 论文内容对象
            output_dir: 输出目录
            formats: 要导出的格式列表 ['html', 'pdf', 'docx', 'md']，默认全部

        Returns:
            格式到文件路径的映射
        """
        if formats is None:
            formats = ['html', 'pdf', 'docx', 'md']

        results = {}

        # 首先生成 HTML（基础格式）
        html_path = output_dir / "article.html"
        html_content = self._generate_html(article_sections, paper_content)
        html_path.write_text(str(html_content), encoding='utf-8')
        results['html'] = html_path
        logger.info(f"HTML 导出成功: {html_path}")

        # 导出 PDF
        if 'pdf' in formats:
            try:
                pdf_path = self._export_pdf(html_path, output_dir)
                if pdf_path and pdf_path.exists():
                    results['pdf'] = pdf_path
                    logger.info(f"PDF 导出成功: {pdf_path}")
                else:
                    logger.warning("PDF 导出失败: 未生成有效文件")
            except Exception as e:
                logger.warning(f"PDF 导出失败: {e}")

        # 导出 Word
        if 'docx' in formats:
            try:
                docx_path = self._export_docx(article_sections, paper_content, output_dir)
                if docx_path and docx_path.exists():
                    results['docx'] = docx_path
                    logger.info(f"Word 导出成功: {docx_path}")
                else:
                    logger.warning("Word 导出失败: 未生成有效文件")
            except Exception as e:
                logger.warning(f"Word 导出失败: {e}")

        # 导出 Markdown
        if 'md' in formats:
            try:
                md_path = self._export_markdown(article_sections, paper_content, output_dir)
                if md_path and md_path.exists():
                    results['md'] = md_path
                    logger.info(f"Markdown 导出成功: {md_path}")
                else:
                    logger.warning("Markdown 导出失败: 未生成有效文件")
            except Exception as e:
                logger.warning(f"Markdown 导出失败: {e}")

        return results

    def _generate_html(self, article_sections, paper_content) -> str:
        """生成完整 HTML（直接调用内部 _build_html，避免临时文件 write 类型问题）"""
        from .renderer import HTMLRenderer
        renderer = HTMLRenderer()
        html_content = renderer._build_html(article_sections, paper_content)
        if not isinstance(html_content, str):
            html_content = str(html_content)
        return html_content

    def _export_pdf(self, html_path: Path, output_dir: Path) -> Path:
        """导出 PDF"""
        from .renderer import PDFExporter
        pdf_path = output_dir / "article.pdf"
        exporter = PDFExporter()
        return exporter.export(html_path, pdf_path)

    def _export_docx(self, article_sections, paper_content, output_dir: Path) -> Path:
        """导出 Word 文档（DOCX）- 包含图片"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise RuntimeError("python-docx 未安装，请运行: pip install python-docx")

        doc = Document()

        # 设置文档样式 - 暖米色主题
        style = doc.styles['Normal']
        style.font.name = 'Noto Sans SC'
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.8
        style.paragraph_format.space_after = Pt(12)

        # 添加标题样式
        title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Noto Serif SC'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(44, 62, 80)  # #2C3E50
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 收集章节信息用于目录
        toc_entries = []
        for section in article_sections:
            if section.section_type not in ['hero', 'paper_info']:
                toc_entries.append(section.title)

        # 处理每个章节
        hero_added = False
        for section in article_sections:
            section_type = section.section_type
            title = section.title
            content = section.content
            image_path = normalize_path(section.image_path) if getattr(section, "image_path", None) else None

            if section_type == "hero":
                self._add_hero_to_docx(doc, title, content, image_path)
                hero_added = True
                # 在Hero区之后添加目录
                if toc_entries:
                    self._add_toc_to_docx(doc, toc_entries)
            elif section_type == "paper_info":
                self._add_paper_info_to_docx(doc, title, content)
            elif section_type == "recommendations":
                self._add_recommendations_to_docx(doc, title, content)
            else:
                self._add_section_to_docx(doc, title, content, image_path)

        # 如果没有hero区但有目录，在最开始添加目录
        if not hero_added and toc_entries:
            self._add_toc_to_docx(doc, toc_entries)

        # 保存文档
        docx_path = output_dir / "article.docx"
        doc.save(str(docx_path))
        return docx_path

    def _add_toc_to_docx(self, doc: 'Document', toc_entries: list):
        """添加目录到 Word 文档"""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not toc_entries:
            return

        # 添加目录标题
        toc_heading = doc.add_paragraph()
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_heading.add_run('目录')
        toc_run.font.name = 'Noto Serif SC'
        toc_run.font.size = Pt(18)
        toc_run.font.bold = True
        toc_run.font.color.rgb = RGBColor(22, 160, 133)
        toc_heading.paragraph_format.space_after = Pt(12)

        for i, entry in enumerate(toc_entries, 1):
            toc_item = doc.add_paragraph()
            toc_item.paragraph_format.left_indent = Inches(0.5)
            toc_item.paragraph_format.space_after = Pt(4)

            num_run = toc_item.add_run(f'{i}. ')
            num_run.font.name = 'Noto Sans SC'
            num_run.font.size = Pt(11)
            num_run.font.color.rgb = RGBColor(22, 160, 133)

            entry_run = toc_item.add_run(entry)
            entry_run.font.name = 'Noto Sans SC'
            entry_run.font.size = Pt(11)

        # 添加分隔线
        doc.add_paragraph()
        separator = doc.add_paragraph('─' * 50)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = separator.runs[0]
        sep_run.font.color.rgb = RGBColor(200, 200, 200)
        doc.add_paragraph()

    def _add_hyperlink(self, paragraph, text: str, url: str):
        """在段落中添加可点击的超链接"""
        from docx.oxml.shared import qn
        from docx.oxml import OxmlElement

        # 创建超链接元素
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), self._create_relationship(paragraph.part, url))

        # 创建 run 元素
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # 设置颜色
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '16A085')  # 绿色
        rPr.append(color)

        # 下划线
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        run.append(t)

        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _create_relationship(self, part, url: str) -> str:
        """创建关系ID用于超链接"""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        return r_id

    def _extract_and_add_links(self, paragraph, text: str):
        """提取文本中的链接并添加为可点击超链接"""
        import re

        # 匹配 [text](url) 格式
        link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'

        last_end = 0
        for match in re.finditer(link_pattern, text):
            # 添加链接前的文本
            if match.start() > last_end:
                pre_text = text[last_end:match.start()]
                if pre_text:
                    run = paragraph.add_run(pre_text)
                    run.font.name = 'Noto Sans SC'
                    run.font.size = Pt(10)

            # 添加超链接
            link_text = match.group(1)
            link_url = match.group(2)
            self._add_hyperlink(paragraph, link_text, link_url)

            last_end = match.end()

        # 添加剩余文本
        if last_end < len(text):
            remaining = text[last_end:]
            if remaining:
                run = paragraph.add_run(remaining)
                run.font.name = 'Noto Sans SC'
                run.font.size = Pt(10)

    def _add_hero_to_docx(self, doc: 'Document', title: str, content: str, image_path: Optional[str]):
        """添加 Hero 区到 Word"""
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        lines = content.split("\n")
        main_title = lines[0].strip() if lines else title

        # 主标题
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(main_title)
        title_run.font.name = 'Noto Serif SC'
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(44, 62, 80)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 副标题
        if len(lines) > 1:
            subtitle = lines[1].strip()
            subtitle_para = doc.add_paragraph()
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.name = 'Noto Sans SC'
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.color.rgb = RGBColor(22, 160, 133)  # #16A085
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        for line in lines[2:]:
            if line.strip() and "**" in line:
                meta_para = doc.add_paragraph()
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 清理 Markdown 标记
                clean_line = line.replace("**", "").strip()
                meta_run = meta_para.add_run(clean_line)
                meta_run.font.size = Pt(10)
                meta_run.font.color.rgb = RGBColor(102, 102, 102)

        # 添加图片
        if image_path and Path(image_path).exists():
            try:
                doc.add_paragraph()  # 空行
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(str(image_path), width=Inches(5))
                doc.add_paragraph()  # 空行
            except Exception as e:
                logger.warning(f"Word 图片添加失败: {e}")

        # 分隔线
        doc.add_paragraph("_" * 60)
        doc.add_paragraph()

    def _add_section_to_docx(self, doc: 'Document', title: str, content: str, image_path: Optional[str]):
        """添加标准章节到 Word - 优化排版"""
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # 章节标题 - 使用更大的字体和更好的样式
        heading = doc.add_heading(level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_run = heading.add_run(title)
        heading_run.font.name = 'Noto Serif SC'
        heading_run.font.size = Pt(20)
        heading_run.font.color.rgb = RGBColor(22, 160, 133)  # 使用主题绿色
        heading_run.font.bold = True
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(10)
        heading.paragraph_format.keep_with_next = True  # 与下一段保持在一起

        # 处理内容 - 清理 Markdown
        clean_content = self._clean_markdown_for_word(content)

        # 分段添加
        paragraphs = clean_content.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # 跳过纯标记行
            if para_text in ['---', '***', '___'] or para_text.startswith('---'):
                continue

            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 1.5  # 适当的行距
            para.paragraph_format.space_after = Pt(8)
            para.paragraph_format.first_line_indent = Inches(0)  # 无首行缩进，使用段间距

            # 处理术语注解和加粗文本
            self._add_formatted_text_to_para(para, para_text)

        # 添加图片 - 更好的布局
        if image_path and Path(image_path).exists():
            try:
                # 在图片前添加分隔空间
                doc.add_paragraph()

                # 图片容器段落
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.paragraph_format.space_before = Pt(12)
                img_para.paragraph_format.space_after = Pt(6)

                run = img_para.add_run()
                # 根据图片比例调整宽度，最大5.5英寸
                run.add_picture(str(image_path), width=Inches(5.5))

                # 图片说明 - 更好的样式
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(12)

                cap_run = caption.add_run(f"▲ 图：{title}")
                cap_run.font.name = 'Noto Sans SC'
                cap_run.font.size = Pt(9)
                cap_run.font.color.rgb = RGBColor(108, 117, 125)  # 灰色

                doc.add_paragraph()  # 空行
            except Exception as e:
                logger.warning(f"Word 图片添加失败: {e}")

    def _add_paper_info_to_docx(self, doc: 'Document', title: str, content: str):
        """添加论文信息到 Word"""
        from docx.shared import Pt, RGBColor

        # 分隔
        doc.add_paragraph()
        doc.add_paragraph("_" * 60)

        # 标题
        heading = doc.add_heading(level=2)
        heading_run = heading.add_run(title)
        heading_run.font.name = 'Noto Serif SC'
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(128, 128, 128)

        # 信息项
        lines = content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 清理 Markdown
            clean_line = line.replace("**", "").strip()
            if clean_line.startswith("-"):
                clean_line = clean_line[1:].strip()

            if clean_line and ":" in clean_line:
                parts = clean_line.split(":", 1)
                label = parts[0].strip()
                value = parts[1].strip() if len(parts) > 1 else ""

                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(6)

                # 标签加粗
                label_run = para.add_run(f"{label}: ")
                label_run.font.name = 'Noto Sans SC'
                label_run.font.bold = True
                label_run.font.size = Pt(10)

                # 值
                if value:
                    value_run = para.add_run(value)
                    value_run.font.name = 'Noto Sans SC'
                    value_run.font.size = Pt(10)

    def _add_recommendations_to_docx(self, doc: 'Document', title: str, content: str):
        """添加推荐章节到 Word - 参考论文信息排版风格"""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        # 调试日志
        logger.debug(f"推荐章节内容长度: {len(content)} 字符")
        logger.debug(f"推荐章节内容预览: {content[:500]}...")

        # 分隔线
        doc.add_paragraph()
        doc.add_paragraph("_" * 60)

        # 标题 - 使用与论文信息相同的样式
        heading = doc.add_heading(level=2)
        heading_run = heading.add_run(title)
        heading_run.font.name = 'Noto Serif SC'
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(128, 128, 128)

        # 处理内容
        lines = content.split('\n')
        processed_count = 0
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            i += 1

            if not line:
                continue

            # 跳过说明文字、分隔线
            if '基于学术论文引用网络' in line or line == '---':
                continue

            # 跳过一键解读链接
            if '一键解读' in line or '📄' in line:
                continue

            # 处理子标题 (### 🔬 相关论文推荐、### 📚 引用网络 等)
            if line.startswith('###'):
                sub_title = line.replace('###', '').strip()
                # 移除所有常见emoji
                sub_title = sub_title.replace('🔬', '').replace('📚', '').replace('🔍', '').replace('💡', '').strip()

                sub_para = doc.add_paragraph()
                sub_run = sub_para.add_run(sub_title)
                sub_run.font.name = 'Noto Serif SC'
                sub_run.font.size = Pt(11)
                sub_run.font.bold = True
                sub_run.font.color.rgb = RGBColor(22, 160, 133)
                sub_para.paragraph_format.space_before = Pt(8)
                sub_para.paragraph_format.space_after = Pt(4)
                continue

            # 处理论文标题（**1. Title** (2024)）
            title_match = re.match(r'\*\*(\d+)\.\s*([^*]+?)\*\*\s*\((\d{4})\)', line)
            if title_match:
                num, paper_title, year = title_match.groups()

                # 论文标题作为一个段落
                title_para = doc.add_paragraph()
                title_para.paragraph_format.space_before = Pt(8)
                title_para.paragraph_format.space_after = Pt(4)

                num_run = title_para.add_run(f"{num}. ")
                num_run.font.name = 'Noto Sans SC'
                num_run.font.size = Pt(10)
                num_run.font.bold = True
                num_run.font.color.rgb = RGBColor(22, 160, 133)

                title_run = title_para.add_run(f"{paper_title} ({year})")
                title_run.font.name = 'Noto Serif SC'
                title_run.font.size = Pt(10)
                title_run.font.bold = True
                continue

            # 处理 **标签**: 值 格式（作者、简介、链接、推荐理由等）
            label_match = re.match(r'\*\*([^*]+)\*\*:\s*(.+)', line)
            if label_match:
                label = label_match.group(1).strip()
                value = label_match.group(2).strip()

                # 跳过一键解读
                if '一键解读' in label:
                    continue

                item_para = doc.add_paragraph()
                item_para.paragraph_format.space_after = Pt(3)
                item_para.paragraph_format.left_indent = Inches(0.2)

                # 标签加粗
                label_run = item_para.add_run(f"{label}: ")
                label_run.font.name = 'Noto Sans SC'
                label_run.font.size = Pt(10)
                label_run.font.bold = True

                # 处理值（可能包含链接）
                link_match = re.search(r'\[([^\]]+)\]\(([^)]+)\)', value)
                if link_match:
                    link_text = link_match.group(1)
                    link_url = link_match.group(2)
                    self._add_hyperlink(item_para, link_text, link_url)
                else:
                    # 普通文本值
                    value_run = item_para.add_run(value)
                    value_run.font.name = 'Noto Sans SC'
                    value_run.font.size = Pt(10)
                continue

            # 处理子标题（如：**引用该论文的研究：**）
            # 使用正则表达式精确匹配只有标题没有值的情况（支持中英文冒号）
            subheading_match = re.match(r'\*\*([^:*]+?)\*\*[:：]\s*$', line)
            if subheading_match:
                sub_heading = subheading_match.group(1).strip()
                sub_para = doc.add_paragraph()
                sub_run = sub_para.add_run(sub_heading)
                sub_run.font.name = 'Noto Sans SC'
                sub_run.font.size = Pt(10)
                sub_run.font.bold = True
                sub_run.font.color.rgb = RGBColor(44, 62, 80)
                sub_para.paragraph_format.space_before = Pt(6)
                sub_para.paragraph_format.space_after = Pt(3)
                continue

            # 处理引用网络列表项 (- [Title](url) 或 - [Title](url) (year))
            if line.startswith('- ['):
                # 支持可选的年份: - [Title](url) 或 - [Title](url) (year)
                link_match = re.search(r'- \[([^\]]+)\]\(([^)]+)\)(?:\s*\((\d{4})\))?', line)
                if link_match:
                    title_text = link_match.group(1)
                    url = link_match.group(2)
                    year = link_match.group(3) if link_match.group(3) else ""

                    item_para = doc.add_paragraph()
                    item_para.paragraph_format.space_after = Pt(3)
                    item_para.paragraph_format.left_indent = Inches(0.3)

                    # 添加圆点符号
                    bullet_run = item_para.add_run("• ")
                    bullet_run.font.name = 'Noto Sans SC'
                    bullet_run.font.size = Pt(10)

                    # 添加超链接
                    display_text = f"{title_text} ({year})" if year else title_text
                    self._add_hyperlink(item_para, display_text, url)
                    processed_count += 1
                else:
                    logger.debug(f"未匹配的引用行: {line[:100]}")
                continue

            # 记录未处理的行（用于调试）
            logger.debug(f"推荐章节未处理的行: {line[:100]}")

        logger.debug(f"推荐章节处理了 {processed_count} 个条目")

    def _clean_markdown_for_word(self, text: str) -> str:
        """清理 Markdown 标记以便 Word 显示 - 完全版本"""
        import re

        # 移除代码块标记
        text = re.sub(r'```\w*\n?', '', text)
        text = re.sub(r'```', '', text)

        # 移除行内代码标记
        text = re.sub(r'`([^`]+)`', r'\1', text)

        # 移除 Markdown 标题标记
        text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)

        # 移除水平分隔线
        text = re.sub(r'\n---\n', '\n', text)
        text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)

        # 处理术语注解 *术语（解释）* -> 保留术语和解释
        text = re.sub(r'\*([^*（]+?)（(.+?)）\*', r'\1（\2）', text)

        # 处理加粗标记 **text** -> 保留文本
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)

        # 处理斜体标记 *text* -> 保留文本
        text = re.sub(r'\*([^*]+)\*', r'\1', text)

        # 处理链接 - [text](url) -> text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

        # 移除表情符号前的标记
        text = re.sub(r'^(###|##|#)\s*', '', text, flags=re.MULTILINE)

        # 处理列表标记 - 简化为普通文本
        lines = text.split('\n')
        result_lines = []
        for line in lines:
            original = line
            line = line.strip()

            # 跳过空行
            if not line:
                result_lines.append('')
                continue

            # 处理列表项
            if line.startswith('- ') or line.startswith('* '):
                line = line[2:].strip()
            elif re.match(r'^\d+\.[\s\u3000]+', line):
                # 数字列表保留数字，去掉多余空格
                line = re.sub(r'^(\d+)\.\s+', r'\1. ', line)

            # 再次处理加粗标记（可能在处理列表后还有剩余）
            line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)

            # 再次处理斜体标记
            line = re.sub(r'\*([^*]+)\*', r'\1', line)

            # 移除前导空格
            line = line.lstrip()

            result_lines.append(line)

        # 合并连续的空行
        final_lines = []
        prev_empty = False
        for line in result_lines:
            is_empty = not line.strip()
            if is_empty and prev_empty:
                continue
            final_lines.append(line)
            prev_empty = is_empty

        return '\n'.join(final_lines)

    def _process_terms_for_word(self, text: str) -> str:
        """处理术语注解为 Word 友好格式"""
        # 将 *术语（解释）* 转换为 "术语（解释）"
        pattern = r'\*([^*（]+?)（(.+?)）\*'

        def replace_term(match):
            term = match.group(1)
            explanation = match.group(2)
            return f"{term}（{explanation}）"

        return re.sub(pattern, replace_term, text)

    def _add_formatted_text_to_para(self, para, text: str):
        """向段落添加格式化文本（支持加粗、术语高亮）"""
        from docx.shared import Pt, RGBColor

        # 处理术语注解 *术语（解释）*
        term_pattern = r'\*([^*（]+?)（(.+?)）\*'

        # 分割文本：普通文本、加粗文本 **text**、术语 *text、*
        parts = []
        last_end = 0

        # 找到所有加粗和术语模式
        patterns = [
            (r'\*\*(.+?)\*\*', 'bold'),  # 加粗
            (term_pattern, 'term'),  # 术语注解
        ]

        for pattern, ptype in patterns:
            for match in re.finditer(pattern, text):
                # 添加之前的普通文本
                if match.start() > last_end:
                    parts.append((text[last_end:match.start()], 'normal'))

                if ptype == 'bold':
                    parts.append((match.group(1), 'bold'))
                elif ptype == 'term':
                    term = match.group(1)
                    explanation = match.group(2)
                    parts.append((f"{term}（{explanation}）", 'term'))

                last_end = match.end()

        # 添加剩余文本
        if last_end < len(text):
            parts.append((text[last_end:], 'normal'))

        # 如果没有找到任何格式，直接添加整体文本
        if not parts:
            parts = [(text, 'normal')]

        # 添加到段落
        for content, style in parts:
            if not content:
                continue
            run = para.add_run(content)
            run.font.name = 'Noto Sans SC'
            run.font.size = Pt(11)

            if style == 'bold':
                run.font.bold = True
                run.font.color.rgb = RGBColor(44, 62, 80)
            elif style == 'term':
                # 术语使用绿色高亮
                run.font.color.rgb = RGBColor(22, 160, 133)
                run.font.bold = True

    def _export_markdown(self, article_sections, paper_content, output_dir: Path) -> Path:
        """导出 Markdown（包含 base64 图片）"""
        md_content = []

        # 添加 YAML frontmatter
        md_content.append("---")
        md_content.append(f'title: "{paper_content.title or "论文解读"}"')
        md_content.append(f'author: "Paper Interpreter"')
        md_content.append(f'date: "{paper_content.publication_date or ""}"')
        md_content.append("---")
        md_content.append("")

        # 处理每个章节
        for section in article_sections:
            section_type = section.section_type
            title = section.title
            content = section.content
            image_path = normalize_path(section.image_path) if getattr(section, "image_path", None) else None

            if section_type == "hero":
                # Hero 区特殊处理
                lines = content.split("\n")
                if lines:
                    md_content.append(f"# {lines[0]}")
                    if len(lines) > 1:
                        md_content.append(f"")
                        md_content.append(f"> {lines[1]}")
                    md_content.append("")
                    # 元信息
                    for line in lines[2:]:
                        if line.strip():
                            clean_line = line.replace("**", "")
                            md_content.append(f"*{clean_line}*")
                    md_content.append("")
            elif section_type == "paper_info":
                md_content.append(f"## {title}")
                md_content.append("")
                lines = content.split("\n")
                for line in lines:
                    if line.strip():
                        clean_line = line.replace("**", "**")
                        md_content.append(clean_line)
                md_content.append("")
            else:
                # 标准章节
                md_content.append(f"## {title}")
                md_content.append("")

                # 处理内容
                clean_content = self._clean_markdown_for_md(content)
                md_content.append(clean_content)
                md_content.append("")

            # 添加图片（base64 嵌入）
            image_path = normalize_path(image_path)
            if image_path and Path(image_path).exists():
                try:
                    base64_img = self._image_to_base64(image_path)
                    if base64_img:
                        ext = Path(image_path).suffix.lower().replace('.', '')
                        if ext == 'jpg':
                            ext = 'jpeg'
                        md_content.append(f"![{title}]({base64_img})")
                        md_content.append(f"*图：{title}*")
                        md_content.append("")
                except Exception as e:
                    logger.warning(f"Markdown 图片嵌入失败: {e}")

        # 保存
        md_path = output_dir / "article.md"
        md_path.write_text('\n'.join(str(c) for c in md_content), encoding='utf-8')
        return md_path

    def _clean_markdown_for_md(self, text: str) -> str:
        """清理 Markdown 内容"""
        # 处理术语注解
        pattern = r'\*([^*（]+?)（(.+?)）\*'

        def replace_term(match):
            term = match.group(1)
            explanation = match.group(2)
            return f"**{term}**（{explanation}）"

        text = re.sub(pattern, replace_term, text)

        # 清理其他标记
        text = re.sub(r'```\w*\n?', '', text)
        text = re.sub(r'```', '', text)

        return text.strip()

    def _image_to_base64(self, image_path: str) -> str:
        """将图片转换为 base64"""
        image_path = normalize_path(image_path)
        if not image_path:
            return ""
        try:
            with open(image_path, "rb") as f:
                image_data = f.read()
            base64_data = base64.b64encode(image_data).decode('utf-8')

            ext = Path(image_path).suffix.lower()
            mime_type = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif': 'image/gif',
                '.webp': 'image/webp'
            }.get(ext, 'image/png')

            return f"data:{mime_type};base64,{base64_data}"
        except Exception as e:
            logger.warning(f"图片转换失败: {e}")
            return ""
